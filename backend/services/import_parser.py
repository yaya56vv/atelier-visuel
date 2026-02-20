"""Service import — Parsing universel de documents uploadés.

Porte d'entrée universelle du stockage énergétique.
Extrait le texte de TOUT type de document pour l'indexation :

- PDF → texte via pdfplumber (fallback PyPDF2)
- DOCX/Word → texte via python-docx
- Texte/Markdown/Code → lecture directe (UTF-8/Latin-1)
- JSON → formaté lisiblement
- CSV → converti en texte tabulé
- YouTube → transcription via youtube-transcript-api
- Audio/Podcast → transcription via Whisper (si disponible)
- Images → pas d'extraction texte (OCR futur)
"""

import os
import json
import re
from pathlib import Path

# Limite globale de texte pour l'indexation
MAX_TEXT_CHARS = 30000


async def parse_uploaded_file(
    file_path: str, content_type: str, original_filename: str
) -> str | None:
    """Extrait le texte d'un fichier uploadé.
    
    Retourne le texte extrait ou None si pas d'extraction possible.
    """
    ext = Path(original_filename).suffix.lower()
    
    # PDF
    if content_type == "pdf" or ext == ".pdf":
        return await _extract_pdf_text(file_path)
    
    # Word / DOCX
    if content_type == "docx" or ext in (".docx", ".doc"):
        return _extract_docx_text(file_path)
    
    # JSON
    if content_type == "json" or ext == ".json":
        return _extract_json(file_path)
    
    # CSV / TSV
    if ext in (".csv", ".tsv"):
        return _extract_csv(file_path)
    
    # Code source et texte
    if content_type == "texte" or content_type == "code" or ext in TEXT_EXTENSIONS:
        return _extract_text_file(file_path)
    
    # YouTube (contenu = URL, pas un fichier physique — géré dans parse_youtube_url)
    # Audio / Podcast → transcription Whisper
    if content_type == "audio" or ext in AUDIO_EXTENSIONS:
        return await _transcribe_audio(file_path, original_filename)
    
    # Images, fichiers génériques → pas d'extraction texte
    return None


# ═══════════════════════════════════════════════════════════
# EXTENSIONS RECONNUES
# ═══════════════════════════════════════════════════════════

TEXT_EXTENSIONS = {
    # Texte brut
    ".txt", ".text",
    # Markdown
    ".md", ".markdown", ".mdx",
    # Code Python
    ".py", ".pyw", ".pyi",
    # Code JavaScript / TypeScript
    ".js", ".jsx", ".ts", ".tsx", ".mjs", ".cjs",
    # Code Web
    ".html", ".htm", ".css", ".scss", ".sass", ".less",
    # Code divers
    ".java", ".c", ".cpp", ".h", ".hpp", ".cs", ".go", ".rs", ".rb",
    ".php", ".swift", ".kt", ".scala", ".r", ".sql", ".sh", ".bash",
    ".ps1", ".bat", ".cmd",
    # Config / Data
    ".yaml", ".yml", ".toml", ".ini", ".cfg", ".conf", ".env",
    ".xml", ".svg",
    # Documentation
    ".rst", ".tex", ".adoc", ".org",
    # Logs
    ".log",
}

AUDIO_EXTENSIONS = {
    ".mp3", ".wav", ".m4a", ".ogg", ".flac", ".aac", ".wma", ".opus",
}


# ═══════════════════════════════════════════════════════════
# EXTRACTEURS PAR TYPE
# ═══════════════════════════════════════════════════════════

async def _extract_pdf_text(file_path: str) -> str | None:
    """Extrait le texte d'un PDF."""
    
    # Essayer pdfplumber d'abord (meilleure qualité)
    try:
        import pdfplumber
        text_parts = []
        with pdfplumber.open(file_path) as pdf:
            for i, page in enumerate(pdf.pages):
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(f"--- Page {i+1} ---\n{page_text}")
        
        if text_parts:
            return _truncate("\n\n".join(text_parts))
    except ImportError:
        pass
    except Exception as e:
        print(f"[import_parser] Erreur pdfplumber: {e}")
    
    # Fallback : PyPDF2
    try:
        from PyPDF2 import PdfReader
        reader = PdfReader(file_path)
        text_parts = []
        for i, page in enumerate(reader.pages):
            page_text = page.extract_text()
            if page_text:
                text_parts.append(f"--- Page {i+1} ---\n{page_text}")
        
        if text_parts:
            return _truncate("\n\n".join(text_parts))
    except ImportError:
        print("[import_parser] Ni pdfplumber ni PyPDF2 installé.")
    except Exception as e:
        print(f"[import_parser] Erreur PyPDF2: {e}")
    
    return None


def _extract_docx_text(file_path: str) -> str | None:
    """Extrait le texte d'un fichier Word (.docx)."""
    try:
        from docx import Document
        doc = Document(file_path)
        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                # Préserver les styles de titre
                if para.style and para.style.name.startswith("Heading"):
                    level = para.style.name.replace("Heading ", "").strip()
                    prefix = "#" * (int(level) if level.isdigit() else 1)
                    paragraphs.append(f"{prefix} {text}")
                else:
                    paragraphs.append(text)
        
        # Extraire aussi les tableaux
        for table in doc.tables:
            table_lines = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                table_lines.append(" | ".join(cells))
            if table_lines:
                paragraphs.append("\n[Tableau]\n" + "\n".join(table_lines))
        
        if paragraphs:
            return _truncate("\n\n".join(paragraphs))
    except ImportError:
        print("[import_parser] python-docx non installé. pip install python-docx")
    except Exception as e:
        print(f"[import_parser] Erreur DOCX: {e}")
    
    return None


def _extract_json(file_path: str) -> str | None:
    """Formate un fichier JSON de manière lisible."""
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
        # Formater joliment pour l'indexation
        formatted = json.dumps(data, indent=2, ensure_ascii=False)
        return _truncate(formatted)
    except Exception as e:
        # Fallback : lire comme texte brut
        print(f"[import_parser] JSON invalide, lecture brute: {e}")
        return _extract_text_file(file_path)


def _extract_csv(file_path: str) -> str | None:
    """Extrait le contenu d'un CSV en texte tabulé."""
    try:
        import csv
        lines = []
        with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
            # Détecter le délimiteur
            sample = f.read(4096)
            f.seek(0)
            try:
                dialect = csv.Sniffer().sniff(sample)
            except csv.Error:
                dialect = csv.excel
            
            reader = csv.reader(f, dialect)
            for i, row in enumerate(reader):
                lines.append(" | ".join(row))
                if i > 500:  # Limiter les très gros CSV
                    lines.append(f"[... {i} lignes affichées sur total]")
                    break
        
        if lines:
            return _truncate("\n".join(lines))
    except Exception as e:
        print(f"[import_parser] Erreur CSV: {e}")
        return _extract_text_file(file_path)
    
    return None


def _extract_text_file(file_path: str) -> str | None:
    """Lit un fichier texte brut, markdown, ou code source."""
    try:
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
        except UnicodeDecodeError:
            with open(file_path, 'r', encoding='latin-1') as f:
                text = f.read()
        
        if not text.strip():
            return None
        
        return _truncate(text)
    except Exception as e:
        print(f"[import_parser] Erreur lecture texte: {e}")
        return None


# ═══════════════════════════════════════════════════════════
# YOUTUBE — Extraction de transcription
# ═══════════════════════════════════════════════════════════

def extract_youtube_id(url: str) -> str | None:
    """Extrait l'ID vidéo d'une URL YouTube."""
    patterns = [
        r'(?:youtube\.com/watch\?v=|youtu\.be/|youtube\.com/embed/)([a-zA-Z0-9_-]{11})',
        r'youtube\.com/shorts/([a-zA-Z0-9_-]{11})',
    ]
    for pattern in patterns:
        match = re.search(pattern, url)
        if match:
            return match.group(1)
    return None


async def parse_youtube_url(url: str) -> dict:
    """Télécharge la transcription d'une vidéo YouTube.
    
    Retourne:
        {
            "video_id": str,
            "title": str | None,
            "transcript": str | None,
            "error": str | None,
        }
    """
    video_id = extract_youtube_id(url)
    if not video_id:
        return {"video_id": None, "transcript": None, "error": "URL YouTube invalide"}
    
    result = {"video_id": video_id, "title": None, "transcript": None, "error": None}
    
    # Récupérer le titre via une requête légère (oEmbed)
    try:
        import httpx
        async with httpx.AsyncClient(timeout=10) as client:
            resp = await client.get(
                f"https://www.youtube.com/oembed?url=https://www.youtube.com/watch?v={video_id}&format=json"
            )
            if resp.status_code == 200:
                data = resp.json()
                result["title"] = data.get("title")
    except Exception:
        pass
    
    # Récupérer la transcription
    try:
        from youtube_transcript_api import YouTubeTranscriptApi
        
        # Essayer français d'abord, puis anglais, puis auto-généré
        transcript_list = None
        for lang_codes in [['fr'], ['en'], ['fr', 'en']]:
            try:
                transcript_list = YouTubeTranscriptApi.get_transcript(video_id, languages=lang_codes)
                break
            except Exception:
                continue
        
        if transcript_list is None:
            # Essayer n'importe quelle langue disponible
            try:
                transcripts = YouTubeTranscriptApi.list_transcripts(video_id)
                # Prendre la première disponible
                for t in transcripts:
                    transcript_list = t.fetch()
                    break
            except Exception:
                pass
        
        if transcript_list:
            # Formater la transcription avec timestamps
            lines = []
            for entry in transcript_list:
                start = int(entry.get('start', 0))
                minutes = start // 60
                seconds = start % 60
                text = entry.get('text', '')
                lines.append(f"[{minutes:02d}:{seconds:02d}] {text}")
            
            result["transcript"] = _truncate("\n".join(lines))
        else:
            result["error"] = "Aucune transcription disponible pour cette vidéo"
    
    except ImportError:
        result["error"] = "youtube-transcript-api non installé"
    except Exception as e:
        result["error"] = f"Erreur transcription: {e}"
    
    return result


# ═══════════════════════════════════════════════════════════
# AUDIO — Transcription via Whisper local
# ═══════════════════════════════════════════════════════════

async def _transcribe_audio(file_path: str, original_filename: str) -> str | None:
    """Transcrit un fichier audio via Whisper local (si disponible).
    
    Utilise le serveur Whisper MCP s'il est actif, ou whisper CLI en fallback.
    """
    import subprocess
    
    # Essayer whisper CLI (installé globalement ou via pip)
    for whisper_cmd in ["whisper", "whisper-ctranslate2"]:
        try:
            result = subprocess.run(
                [whisper_cmd, file_path,
                 "--model", "large-v3",
                 "--language", "fr",
                 "--output_format", "txt",
                 "--output_dir", str(Path(file_path).parent)],
                capture_output=True, text=True, timeout=300
            )
            
            if result.returncode == 0:
                # Lire le fichier .txt généré
                txt_path = Path(file_path).with_suffix(".txt")
                if txt_path.exists():
                    text = txt_path.read_text(encoding='utf-8')
                    # Nettoyer le fichier temporaire
                    txt_path.unlink(missing_ok=True)
                    if text.strip():
                        return _truncate(f"[Transcription audio : {original_filename}]\n\n{text}")
        except (FileNotFoundError, subprocess.TimeoutExpired):
            continue
        except Exception as e:
            print(f"[import_parser] Erreur Whisper ({whisper_cmd}): {e}")
    
    # Whisper non disponible — signaler à l'utilisateur
    print(f"[import_parser] Whisper non disponible pour transcrire: {original_filename}")
    return None


# ═══════════════════════════════════════════════════════════
# UTILITAIRES
# ═══════════════════════════════════════════════════════════

def _truncate(text: str) -> str:
    """Tronque le texte à MAX_TEXT_CHARS avec indication."""
    if len(text) > MAX_TEXT_CHARS:
        return text[:MAX_TEXT_CHARS] + f"\n\n[... tronqué à {MAX_TEXT_CHARS} caractères sur {len(text)} total]"
    return text
